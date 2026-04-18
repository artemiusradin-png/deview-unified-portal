from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "TE-Credit-Review-and-Borrower-Decision-Portal.docx"


base_features = [
    "Secure internal staff login and role-based access control",
    "Borrower search and profile page",
    "Automatic TE Credit data refresh once per day, subject to approved access method",
    "Storage of latest permitted TE Credit borrower data on Supabase or private server database",
    "TE Credit review module for lending amount, period, approval / rejection, late payment, default, past transactions, bad debt write-off, interest, repayment condition, and arrears indicators where available",
    "Last sync date/time and sync status display",
    "AI-generated borrower summary",
    "AI-assisted good / bad borrower grade with explanation",
    "Basic risk labels for late payment, default, bad debt write-off, arrears, and repeated rejection history",
    "Inquiry and access audit log, basic admin user management, and manual refresh for authorised staff",
]

ops_features = [
    "Advanced search filters by borrower grade, risk status, late payment, default, write-off, approval / rejection, and last refresh date",
    "Borrower review dashboard for recent checks and high-risk cases",
    "Inquiry cost tracking for TE Credit checks",
    "Export / print-ready borrower review report",
    "Case notes for internal staff review",
    "Document / screenshot attachment support for TE Credit evidence",
    "Sync failure alerts for admin users",
    "Enhanced admin permissions by role",
    "UAT support, launch setup, and staff onboarding",
]

management_features = [
    "Management dashboard with borrower grade distribution, risk trends, and review activity",
    "AI-assisted approval / rejection reasoning template for internal use",
    "Borrower trend interpretation for borrowing, repayment, and arrears movement where TE Credit data supports it",
    "Risk priority scoring for case review order",
    "Reminder / alert module for borrowers with status changes or high-risk indicators",
    "Team workflow notes and review status tracking",
    "Monthly management report export",
    "Priority post-launch optimization period after go-live",
]

matrix_rows = [
    ("Secure internal web portal", True, True, True),
    ("Staff login and access control", True, True, True),
    ("Borrower search and profile page", True, True, True),
    ("Daily automatic TE Credit data refresh", True, True, True),
    ("Supabase or private server data storage", True, True, True),
    ("TE Credit borrower data module", True, True, True),
    ("Last sync status and refresh timestamp", True, True, True),
    ("AI borrower summary", True, True, True),
    ("AI-assisted good / bad borrower grade", True, True, True),
    ("AI explanation of borrower grade", True, True, True),
    ("Basic risk labels", True, True, True),
    ("Inquiry and access audit log", True, True, True),
    ("Manual refresh for authorised staff", True, True, True),
    ("Advanced risk and grade filters", False, True, True),
    ("Operations dashboard", False, True, True),
    ("TE Credit inquiry cost tracking", False, True, True),
    ("Export / print borrower report", False, True, True),
    ("Internal case notes", False, True, True),
    ("Document / screenshot attachments", False, True, True),
    ("Sync failure alerts", False, True, True),
    ("Enhanced role permissions", False, True, True),
    ("UAT, launch setup, and staff onboarding", False, True, True),
    ("Management dashboard", False, False, True),
    ("Borrower risk trend interpretation", False, False, True),
    ("AI approval / rejection reasoning template", False, False, True),
    ("Risk priority scoring", False, False, True),
    ("Reminder / alert module", False, False, True),
    ("Review status workflow", False, False, True),
    ("Monthly management report export", False, False, True),
    ("Priority post-launch optimization", False, False, True),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Title", 26), ("Heading 1", 16), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(15, 23, 42)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(7)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(cell, items, size=8.3):
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.18)
        p.paragraph_format.space_after = Pt(1.8)
        run = p.add_run(item)
        run.font.size = Pt(size)


def add_page_break(doc):
    doc.add_page_break()


def build_doc():
    doc = Document()
    style_doc(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.add_run("TE Credit Review & Borrower Decision Portal").bold = True
    title.paragraph_format.space_after = Pt(12)

    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(["Prepared for: Mr. Lam", "Prepared by: DeView", "Date: April 16, 2026"]):
        cell = meta.cell(0, i)
        set_cell_shading(cell, "F8FAFC")
        set_cell_text(cell, text, bold=True, size=9)
        set_cell_width(cell, 6.0)

    doc.add_paragraph()
    add_heading(doc, "Project Overview")
    overview = [
        "DeView proposes to build a secure internal web portal for Eden's Father that works around the client's existing TE Credit workflow. The portal will automatically refresh permitted TE Credit borrower information once per day, organize the data into a clear review interface, and support staff with AI-assisted borrower summaries and good / bad grade indicators.",
        "The objective is not to replace TE Credit. TE Credit remains the credit-reference source. The DeView portal becomes the client's internal operating layer for search, review, decision support, history tracking, AI summaries, audit logs, and management visibility.",
        "Automatic TE Credit access will be implemented through the approved method available to the client, such as API, export, or permitted automated access. All access should remain limited to authorised borrower inquiries and subject to TE Credit's terms and Hong Kong personal data requirements.",
    ]
    for text in overview:
        add_para(doc, text)

    add_page_break(doc)

    pricing_section = doc.add_section(WD_ORIENT.LANDSCAPE)
    pricing_section.top_margin = Cm(1.05)
    pricing_section.bottom_margin = Cm(1.05)
    pricing_section.left_margin = Cm(1.1)
    pricing_section.right_margin = Cm(1.1)
    add_heading(doc, "Pricing Options")

    pricing = doc.add_table(rows=1, cols=3)
    pricing.alignment = WD_TABLE_ALIGNMENT.CENTER
    pricing.autofit = False
    packages = [
        ("Option 1: TE Credit AI Portal", "USD 4,000", "This package includes the essential portal plus automatic daily TE Credit data refresh and AI-assisted borrower decision support.", base_features, "Business benefit: TE Credit data is refreshed daily, staff can review borrowers in one portal, and AI helps identify whether a borrower appears good, risky, or unsuitable based on available TE Credit information.\n\nImportant note: The AI grade is a decision-support indicator. Final approval or rejection should remain with authorised staff and management."),
        ("Option 2: Operations Control Package", "USD 6,000", "This package includes everything in Option 1, plus stronger operational controls for daily lending review.", ops_features, "Business benefit: This option is best for a live operating team that needs control, traceability, and smoother daily workflow. It helps management see which borrowers are being checked, which cases are risky, and whether TE Credit refreshes are running correctly."),
        ("Option 3: Management Intelligence Package", "USD 8,000", "This package includes everything in Option 2, plus management-level reporting, enhanced AI review, and stronger decision intelligence.", management_features, "Business benefit: This option gives owners and managers better visibility into borrower risk, team activity, and portfolio direction. It turns the portal from a daily lookup tool into a management decision-support system."),
    ]
    for idx, package in enumerate(packages):
        cell = pricing.cell(0, idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_width(cell, 8.9)
        set_cell_shading(cell, "FFFFFF")
        title, price, desc, features, benefit = package
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_after = Pt(4)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(price)
        r2.bold = True
        r2.font.size = Pt(20)
        r2.font.color.rgb = RGBColor(15, 23, 42)
        p2.paragraph_format.space_after = Pt(5)
        p3 = cell.add_paragraph(desc)
        p3.paragraph_format.space_after = Pt(4)
        p3.runs[0].font.size = Pt(8.4)
        add_bullets(cell, features)
        for part in benefit.split("\n\n"):
            p4 = cell.add_paragraph()
            if part.startswith("Business benefit:"):
                r = p4.add_run("Business benefit:")
                r.bold = True
                r.font.size = Pt(8.2)
                rest = part.replace("Business benefit:", "", 1)
                p4.add_run(rest).font.size = Pt(8.2)
            elif part.startswith("Important note:"):
                r = p4.add_run("Important note:")
                r.bold = True
                r.font.size = Pt(8.2)
                rest = part.replace("Important note:", "", 1)
                p4.add_run(rest).font.size = Pt(8.2)
            p4.paragraph_format.space_after = Pt(1.5)

    add_page_break(doc)

    matrix_section = doc.add_section(WD_ORIENT.PORTRAIT)
    matrix_section.top_margin = Cm(0.85)
    matrix_section.bottom_margin = Cm(0.85)
    matrix_section.left_margin = Cm(1.2)
    matrix_section.right_margin = Cm(1.2)
    add_heading(doc, "Feature Comparison")
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Function / Deliverable", "USD 4,000", "USD 6,000", "USD 8,000"]
    widths = [10.2, 2.5, 2.5, 2.5]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E2E8F0")
        set_cell_width(cell, widths[i])
        set_cell_text(cell, header, bold=True, size=8.4, align=WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT)
    for name, a, b, c in matrix_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, size=7.6)
        set_cell_width(cells[0], widths[0])
        for i, value in enumerate([a, b, c], start=1):
            set_cell_width(cells[i], widths[i])
            set_cell_text(cells[i], "✓" if value else "", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Recommended Positioning")
    pos = doc.add_table(rows=1, cols=3)
    pos.autofit = False
    cards = [
        ("HKD 30,000", "Strong launch package with daily TE Credit sync and AI borrower grading"),
        ("HKD 45,000", "Best-value operations package with controls, reports, and team workflow"),
        ("HKD 60,000", "Premium management package with intelligence dashboards and advanced decision support"),
    ]
    for i, (price, text) in enumerate(cards):
        cell = pos.cell(0, i)
        set_cell_width(cell, 5.7)
        set_cell_shading(cell, "F8FAFC")
        p = cell.paragraphs[0]
        r = p.add_run(price)
        r.bold = True
        r.font.size = Pt(14)
        p2 = cell.add_paragraph(text)
        p2.runs[0].font.size = Pt(8.8)
    add_para(doc, "This structure keeps the base package attractive and useful while giving the client clear reasons to upgrade if they want stronger internal control, management reporting, and advanced review workflow.")

    add_page_break(doc)

    final_section = doc.add_section(WD_ORIENT.PORTRAIT)
    final_section.top_margin = Cm(1.2)
    final_section.bottom_margin = Cm(1.2)
    final_section.left_margin = Cm(1.35)
    final_section.right_margin = Cm(1.35)
    add_heading(doc, "Hosting, Compliance & Next Steps")
    two = doc.add_table(rows=1, cols=2)
    two.autofit = False
    for cell in two.row_cells(0):
        set_cell_width(cell, 8.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    left, right = two.row_cells(0)
    set_cell_shading(left, "F8FAFC")
    set_cell_shading(right, "F8FAFC")
    set_cell_text(left, "Hosting, Server, and Maintenance", bold=True, size=12)
    p = left.add_paragraph()
    r = p.add_run("HKD 4,000 / month")
    r.bold = True
    r.font.size = Pt(18)
    p2 = left.add_paragraph("Separate from the one-time build fee, DeView can provide private hosting infrastructure, daily sync maintenance, and technical support. Included services:")
    p2.runs[0].font.size = Pt(9)
    add_bullets(left, [
        "Personal server / cloud hosting setup",
        "Deployment and environment management",
        "Daily TE Credit sync monitoring",
        "Backups",
        "Security updates",
        "Minor bug fixes and maintenance support",
        "Sync failure support",
        "Basic uptime and operational support",
    ], size=9)
    p3 = left.add_paragraph("This monthly service ensures the client does not need to manage technical infrastructure internally and gives them a stable long-term operating environment for the portal and daily TE Credit data refresh.")
    p3.runs[0].font.size = Pt(9)

    set_cell_text(right, "Compliance Note", bold=True, size=12)
    p = right.add_paragraph("The portal will support automatic TE Credit data refresh only through an approved access method available to the client. AI outputs will be used for decision support only and should not replace final human approval.")
    p.runs[0].font.size = Pt(9.5)
    p = right.add_paragraph()
    r = p.add_run("Next Step")
    r.bold = True
    r.font.size = Pt(12)
    p = right.add_paragraph("Once the preferred package is confirmed, DeView can finalize the client-facing proposal with delivery timeline, implementation assumptions, payment schedule, and TE Credit access requirements.")
    p.runs[0].font.size = Pt(9.5)

    steps = doc.add_table(rows=1, cols=4)
    steps.autofit = False
    step_cards = [
        ("01", "Confirm preferred package", "Select from USD 4,000 / 6,000 / 8,000 options"),
        ("02", "Finalize proposal", "Delivery timeline, implementation assumptions, and payment schedule"),
        ("03", "Confirm TE Credit access", "Verify permitted access method and data-retention rules"),
        ("04", "Begin implementation", "DeView commences portal build and integration"),
    ]
    for i, (no, title, text) in enumerate(step_cards):
        cell = steps.cell(0, i)
        set_cell_width(cell, 4.25)
        set_cell_shading(cell, "FFFFFF")
        p = cell.paragraphs[0]
        r = p.add_run(no)
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(51, 65, 85)
        p = cell.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        p = cell.add_paragraph(text)
        p.runs[0].font.size = Pt(8.5)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
